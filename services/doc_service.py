from docx import Document

def export_to_word(transcription: str, analysis: dict, output_path: str):
    doc = Document()
    doc.add_heading("Portuguese Learning Result", level=1)

    doc.add_heading("1. Portuguese Transcription", level=2)
    doc.add_paragraph(transcription)

    doc.add_heading("2. Chinese Translation", level=2)
    doc.add_paragraph(analysis.get("translation_zh", ""))

    doc.add_heading("3. Sentence Explanations", level=2)
    for item in analysis.get("sentence_explanations", []):
        doc.add_paragraph(f"原句：{item.get('sentence', '')}")
        doc.add_paragraph(f"解释：{item.get('explanation_zh', '')}")

    doc.add_heading("4. Key Words", level=2)
    for item in analysis.get("key_words", []):
        doc.add_paragraph(
            f"{item.get('word', '')} | {item.get('meaning_zh', '')} | "
            f"{item.get('pos', '')} | {item.get('example', '')}"
        )

    doc.add_heading("5. Grammar Notes", level=2)
    for note in analysis.get("grammar_notes", []):
        doc.add_paragraph(note)

    doc.save(output_path)